#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
CHAT ACADÉMICO - VERSIÓN DEFINITIVA PARA KALM OS
"""

import sys
import os
import threading
import time
import subprocess
from pathlib import Path

# Instalar dependencias
def instalar_dependencias():
    for pkg in ["flask", "flask-cors", "requests", "python-docx"]:
        try:
            __import__(pkg.replace("-", "_"))
        except ImportError:
            print(f"📦 Instalando {pkg}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

instalar_dependencias()

from flask import Flask, request, render_template_string, jsonify, send_file
from flask_cors import CORS
import requests

HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head><title>Chat Académico</title>
<style>
    body { font-family: Arial, sans-serif; margin: 30px; background: #f4f4f9; }
    .container { max-width: 900px; margin: auto; background: white; padding: 25px; border-radius: 10px; box-shadow: 0 0 20px rgba(0,0,0,0.1); }
    h1 { color: #2c3e50; text-align: center; }
    .controls { display: flex; flex-wrap: wrap; gap: 10px; margin-bottom: 20px; align-items: center; }
    .controls label { font-weight: bold; }
    .controls select, .controls input { padding: 8px; border-radius: 5px; border: 1px solid #ccc; }
    .controls button { padding: 8px 16px; background: #3498db; color: white; border: none; border-radius: 5px; cursor: pointer; }
    .controls button:hover { background: #2980b9; }
    #resultado { background: #f9f9f9; padding: 15px; border-radius: 8px; min-height: 400px; white-space: pre-wrap; font-family: 'Segoe UI', sans-serif; border: 1px solid #ddd; overflow-y: auto; max-height: 600px; }
    .btn-export { margin-top: 10px; display: flex; gap: 10px; }
    .btn-export button { padding: 8px 20px; background: #27ae60; color: white; border: none; border-radius: 5px; cursor: pointer; }
    .btn-export button:hover { background: #1e8449; }
    .estado { margin-top: 10px; color: #888; font-style: italic; }
</style>
</head>
<body>
<div class="container">
    <h1>🤖 Chat Académico</h1>
    <p><small>Sin API Key · Pollinations AI</small></p>
    
    <div class="controls">
        <label>Modelo:</label>
        <select id="modelo">
            <option value="openai">openai</option>
            <option value="claude">claude</option>
            <option value="gemini">gemini</option>
            <option value="deepseek">deepseek</option>
            <option value="llama">llama</option>
            <option value="mistral">mistral</option>
        </select>
        <label>Tema:</label>
        <input type="text" id="tema" placeholder="Ej: Inteligencia Artificial" style="flex:1; min-width:200px;">
        <button onclick="generar()">📝 Generar</button>
        <button onclick="probar()">🔍 Probar</button>
    </div>
    
    <div id="resultado">Bienvenido. Escribe un tema y pulsa "Generar".</div>
    
    <div class="btn-export">
        <button onclick="exportar('txt')">📄 Exportar TXT</button>
        <button onclick="exportar('docx')">📤 Exportar Word</button>
        <button onclick="limpiar()">🗑️ Limpiar</button>
    </div>
    <div id="estado" class="estado"></div>
</div>

<script>
function setEstado(msg) { document.getElementById('estado').innerText = msg; }

async function generar() {
    const tema = document.getElementById('tema').value.trim();
    if (!tema) { alert('Escribe un tema.'); return; }
    const modelo = document.getElementById('modelo').value;
    setEstado('⏳ Generando...');
    try {
        const resp = await fetch('/generar', {
            method: 'POST',
            headers: {'Content-Type': 'application/json'},
            body: JSON.stringify({ tema, modelo })
        });
        const data = await resp.json();
        if (data.error) {
            document.getElementById('resultado').innerHTML = '❌ ' + data.error;
        } else {
            document.getElementById('resultado').innerHTML = data.texto.replace(/\n/g, '<br>');
        }
        setEstado('✅ Listo');
    } catch(e) {
        document.getElementById('resultado').innerHTML = '❌ Error: ' + e.message;
        setEstado('❌ Error');
    }
}

async function probar() {
    const modelo = document.getElementById('modelo').value;
    setEstado('🧪 Probando...');
    try {
        const resp = await fetch('/probar', {
            method: 'POST',
            headers: {'Content-Type': 'application/json'},
            body: JSON.stringify({ modelo })
        });
        const data = await resp.json();
        document.getElementById('resultado').innerHTML = '✅ ' + (data.respuesta || 'OK');
        setEstado('✅ Modelo responde');
    } catch(e) {
        document.getElementById('resultado').innerHTML = '❌ Error: ' + e.message;
        setEstado('❌ Error');
    }
}

async function exportar(formato) {
    const contenido = document.getElementById('resultado').innerText;
    if (!contenido || contenido === 'Bienvenido. Escribe un tema y pulsa "Generar".') {
        alert('No hay contenido para exportar.');
        return;
    }
    setEstado('⏳ Exportando...');
    try {
        const resp = await fetch('/exportar', {
            method: 'POST',
            headers: {'Content-Type': 'application/json'},
            body: JSON.stringify({ formato, contenido })
        });
        const blob = await resp.blob();
        const url = URL.createObjectURL(blob);
        const a = document.createElement('a');
        a.href = url;
        a.download = `trabajo.${formato === 'txt' ? 'txt' : 'docx'}`;
        a.click();
        setEstado('✅ Exportado');
    } catch(e) {
        setEstado('❌ Error al exportar');
        alert('Error: ' + e.message);
    }
}

function limpiar() {
    document.getElementById('resultado').innerHTML = 'Conversación limpiada.';
    setEstado('');
}
</script>
</body>
</html>
"""

class PollinationsClient:
    BASE_URL = "https://text.pollinations.ai/openai"
    
    def chat_completion(self, model, messages):
        payload = {"model": model, "messages": messages, "stream": False}
        response = requests.post(self.BASE_URL, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]

app = Flask(__name__)
CORS(app)
client = PollinationsClient()

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/generar', methods=['POST'])
def generar():
    data = request.json
    tema = data.get('tema', '').strip()
    modelo = data.get('modelo', 'openai')
    if not tema:
        return jsonify({'error': 'Tema vacío'}), 400
    
    prompt = f"""
    Escribe un trabajo académico completo sobre: {tema}
    
    Estructura:
    ## INTRODUCCIÓN
    ## INTRODUCTION (in English)
    ## DESARROLLO
    ## CONCLUSIONES
    ## BIBLIOGRAFÍA
    """
    try:
        texto = client.chat_completion(modelo, [{"role": "user", "content": prompt}])
        return jsonify({'texto': texto})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/probar', methods=['POST'])
def probar():
    data = request.json
    modelo = data.get('modelo', 'openai')
    try:
        resp = client.chat_completion(modelo, [{"role": "user", "content": "Responde OK"}])
        return jsonify({'respuesta': resp})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/exportar', methods=['POST'])
def exportar():
    data = request.json
    formato = data.get('formato', 'txt')
    contenido = data.get('contenido', '')
    if not contenido:
        return jsonify({'error': 'Sin contenido'}), 400
    
    from io import BytesIO
    buffer = BytesIO()
    
    if formato == 'txt':
        buffer.write(contenido.encode('utf-8'))
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='trabajo.txt', mimetype='text/plain')
    else:
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Trabajo Académico', 0)
            for linea in contenido.split('\n'):
                if linea.strip():
                    doc.add_paragraph(linea)
            doc.save(buffer)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name='trabajo.docx', 
                           mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        except ImportError:
            return jsonify({'error': 'python-docx no instalado'}), 500

if __name__ == "__main__":
    port = int(os.environ.get('PORT', 5000))
    
    if "--kalm" in sys.argv:
        print(f"🚀 Chat Académico iniciado en puerto {port}")
        print(f"http://localhost:{port}")
        sys.stdout.flush()
        
        def run():
            app.run(host='0.0.0.0', port=port, debug=False, use_reloader=False)
        
        thread = threading.Thread(target=run, daemon=True)
        thread.start()
        
        try:
            while True:
                time.sleep(1)
        except KeyboardInterrupt:
            sys.exit(0)
    else:
        app.run(host='0.0.0.0', port=port, debug=True)
