#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
CHAT ACADÉMICO INTELIGENTE
- Detecta automáticamente el entorno (Windows GUI o servidor web)
- Usa Pollinations AI (sin API key) como principal
- Fallback a G4F si Pollinations falla
- Interfaz Tkinter en Windows, Flask/HTML en web
- Soporte para Kalm OS: modo --kalm para abrir navegador interno
"""

import sys
import os
import platform
import subprocess
import importlib
import threading
import json
import re
import datetime
import time
from pathlib import Path

# ============================================================
# 1. DETECCIÓN DE ENTORNO
# ============================================================
def es_entorno_grafico():
    """Determina si hay entorno gráfico disponible (Windows o X11)"""
    sistema = platform.system()
    if sistema == "Windows":
        return True
    if "DISPLAY" in os.environ:
        return True
    return False

def es_windows():
    return platform.system() == "Windows"

def es_headless():
    return not es_entorno_grafico()

# ============================================================
# 2. INSTALADOR AUTOMÁTICO DE DEPENDENCIAS
# ============================================================
def instalar_dependencias():
    """Instala automáticamente las dependencias necesarias."""
    dependencias = ["requests"]
    if es_entorno_grafico():
        dependencias.append("python-docx")
        try:
            import tkinter
        except ImportError:
            print("⚠️ Tkinter no encontrado.")
            if es_windows():
                print("Reinstala Python con la opción 'tcl/tk'.")
            else:
                subprocess.run(["sudo", "apt-get", "install", "-y", "python3-tk"], check=False)
    else:
        dependencias.append("flask")
        dependencias.append("flask-cors")
        dependencias.append("python-docx")
    
    for pkg in dependencias:
        try:
            importlib.import_module(pkg.replace("-", "_"))
        except ImportError:
            print(f"📦 Instalando {pkg}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

# Ejecutar instalación al inicio
instalar_dependencias()

# ============================================================
# 3. IMPORTACIÓN DE MÓDULOS
# ============================================================
try:
    import requests
except ImportError:
    requests = None

if es_entorno_grafico():
    try:
        import tkinter as tk
        from tkinter import ttk, scrolledtext, messagebox, filedialog
    except ImportError:
        es_entorno_grafico = lambda: False

if es_headless():
    try:
        from flask import Flask, request, render_template_string, jsonify, send_file
        from flask_cors import CORS
    except ImportError:
        Flask = None
        CORS = None
        print("⚠️ Flask no instalado. Instalando...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "flask flask-cors"])
        from flask import Flask, request, render_template_string, jsonify, send_file
        from flask_cors import CORS

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None

try:
    import g4f
    from g4f.client import Client
    G4F_AVAILABLE = True
except ImportError:
    G4F_AVAILABLE = False
    Client = None

# ============================================================
# 4. CLIENTE POLLINATIONS AI
# ============================================================
class PollinationsClient:
    BASE_URL = "https://text.pollinations.ai/openai"
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({"Content-Type": "application/json"})
    
    def chat_completion(self, model, messages):
        payload = {
            "model": model,
            "messages": messages,
            "stream": False
        }
        try:
            response = self.session.post(self.BASE_URL, json=payload, timeout=60)
            response.raise_for_status()
            data = response.json()
            if "choices" in data and len(data["choices"]) > 0:
                return data["choices"][0]["message"]["content"]
            else:
                raise Exception("Respuesta inesperada de Pollinations")
        except Exception as e:
            raise Exception(f"Error en Pollinations: {str(e)}")

# ============================================================
# 5. CLASE PRINCIPAL
# ============================================================
class ChatAcademicoCore:
    def __init__(self):
        self.pollinations = PollinationsClient() if requests else None
        self.g4f_client = Client() if G4F_AVAILABLE and Client else None
    
    def consultar_ia(self, modelo, mensajes, proveedor="pollinations"):
        if proveedor == "pollinations":
            if self.pollinations is None:
                raise Exception("Pollinations no disponible")
            return self.pollinations.chat_completion(modelo, mensajes)
        elif proveedor == "g4f":
            if self.g4f_client is None:
                raise Exception("G4F no disponible")
            response = self.g4f_client.chat.completions.create(
                model=modelo,
                messages=mensajes,
            )
            return response.choices[0].message.content
        else:
            raise Exception(f"Proveedor desconocido: {proveedor}")
    
    def generar_trabajo(self, tema, modelo, proveedor="pollinations"):
        prompt = f"""
        Escribe un trabajo académico completo sobre el siguiente tema:

        TEMA: {tema}

        La estructura debe ser EXACTAMENTE esta:

        ## INTRODUCCIÓN (en español)
        [Texto de la introducción en español]

        ## INTRODUCTION (in English)
        [Texto de la introducción en inglés]

        ## DESARROLLO
        [Texto del desarrollo, con subtemas y explicaciones]

        ## CONCLUSIONES
        [Texto de las conclusiones]

        ## BIBLIOGRAFÍA
        [Lista de referencias bibliográficas en formato APA]

        Requisitos:
        - El trabajo debe ser completo, bien estructurado y coherente.
        - Cada sección debe estar claramente marcada con los títulos indicados.
        - La introducción en español e inglés deben tener contenido similar pero no idéntico.
        - El desarrollo debe tener al menos 3 subtemas.
        - Las conclusiones deben ser reflexivas y basadas en el desarrollo.
        - La bibliografía debe incluir al menos 5 fuentes relevantes.
        """
        return self.consultar_ia(modelo, [{"role": "user", "content": prompt}], proveedor)

# ============================================================
# 6. MODO WEB (Flask) - para Render / headless
# ============================================================
def ejecutar_web(kalm_mode=False):
    """Arranca el servidor Flask con interfaz HTML.
       Si kalm_mode=True, arranca en segundo plano y devuelve la URL."""
    if Flask is None:
        print("❌ Flask no instalado. Instalando...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "flask flask-cors"])
        from flask import Flask, request, render_template_string, jsonify, send_file
        from flask_cors import CORS
    
    app = Flask(__name__)
    CORS(app)
    core = ChatAcademicoCore()
    
    HTML_TEMPLATE = """
    <!DOCTYPE html>
    <html>
    <head><title>Chat Académico</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 30px; background: #f4f4f9; }
        .container { max-width: 900px; margin: auto; background: white; padding: 25px; border-radius: 10px; box-shadow: 0 0 20px rgba(0,0,0,0.1); }
        h1 { color: #2c3e50; text-align: center; }
        .controls { display: flex; flex-wrap: wrap; gap: 10px; margin-bottom: 20px; align-items: center; }
        .controls label { font-weight: bold; }
        .controls select, .controls input { padding: 8px; border-radius: 5px; border: 1px solid #ccc; }
        .controls button { padding: 8px 16px; background: #3498db; color: white; border: none; border-radius: 5px; cursor: pointer; }
        .controls button:hover { background: #2980b9; }
        #resultado { background: #f9f9f9; padding: 15px; border-radius: 8px; min-height: 400px; white-space: pre-wrap; font-family: 'Segoe UI', sans-serif; border: 1px solid #ddd; overflow-y: auto; max-height: 600px; }
        .btn-export { margin-top: 10px; display: flex; gap: 10px; }
        .btn-export button { padding: 8px 20px; background: #27ae60; color: white; border: none; border-radius: 5px; cursor: pointer; }
        .btn-export button:hover { background: #1e8449; }
        .estado { margin-top: 10px; color: #888; font-style: italic; }
        .titulo { color: #c0392b; font-weight: bold; font-size: 1.2em; }
        .subtitulo { color: #d35400; font-weight: bold; }
        .usuario { color: #2980b9; font-weight: bold; }
        .ia { color: #27ae60; }
        hr { border: 1px solid #ddd; }
    </style>
    </head>
    <body>
    <div class="container">
        <h1>🤖 Chat Académico</h1>
        <p><small>Sin API Key · Pollinations AI · Genera trabajos con Introducción (ES/EN), Desarrollo, Conclusiones y Bibliografía.</small></p>
        
        <div class="controls">
            <label>Proveedor:</label>
            <select id="proveedor">
                <option value="pollinations">Pollinations</option>
                <option value="g4f">G4F</option>
            </select>
            
            <label>Modelo:</label>
            <select id="modelo">
                <option value="openai">openai</option>
                <option value="claude">claude</option>
                <option value="gemini">gemini</option>
                <option value="deepseek">deepseek</option>
                <option value="llama">llama</option>
                <option value="mistral">mistral</option>
                <option value="phi">phi</option>
                <option value="qwen">qwen</option>
            </select>
            
            <label>Tema:</label>
            <input type="text" id="tema" placeholder="Ej: Inteligencia Artificial en la educación" style="flex:1; min-width:200px;">
            <button onclick="generar()">📝 Generar Trabajo</button>
            <button onclick="probar()">🔍 Probar</button>
        </div>
        
        <div id="resultado">Bienvenido. Escribe un tema y pulsa "Generar Trabajo".</div>
        
        <div class="btn-export">
            <button onclick="exportar('txt')">📄 Exportar TXT</button>
            <button onclick="exportar('docx')">📤 Exportar Word</button>
            <button onclick="limpiar()">🗑️ Limpiar</button>
        </div>
        <div id="estado" class="estado"></div>
    </div>
    
    <script>
        function actualizarModelos() {
            const prov = document.getElementById('proveedor').value;
            const modelos = {
                'pollinations': ['openai','claude','gemini','deepseek','llama','mistral','phi','qwen'],
                'g4f': ['gpt-4o-mini','gpt-4o','gpt-3.5-turbo','claude-3-haiku','gemini-pro','llama-3-70b']
            };
            const sel = document.getElementById('modelo');
            sel.innerHTML = '';
            modelos[prov].forEach(m => {
                const opt = document.createElement('option');
                opt.value = m;
                opt.text = m;
                sel.appendChild(opt);
            });
        }
        document.getElementById('proveedor').addEventListener('change', actualizarModelos);
        actualizarModelos();
        
        function setEstado(msg) {
            document.getElementById('estado').innerText = msg;
        }
        
        async function generar() {
            const tema = document.getElementById('tema').value.trim();
            if (!tema) { alert('Escribe un tema.'); return; }
            const prov = document.getElementById('proveedor').value;
            const modelo = document.getElementById('modelo').value;
            setEstado('⏳ Generando...');
            try {
                const resp = await fetch('/generar', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({ tema, modelo, proveedor: prov })
                });
                const data = await resp.json();
                if (data.error) {
                    document.getElementById('resultado').innerHTML = '❌ ' + data.error;
                } else {
                    let html = data.texto.replace(/##\s+(.*)/g, '<div class="titulo">📌 $1</div>');
                    html = html.replace(/\*\*\*(.*)\*\*\*/g, '<div class="subtitulo">$1</div>');
                    html = html.replace(/\*\*(.*)\*\*/g, '<div class="subtitulo">$1</div>');
                    html = html.replace(/\n/g, '<br>');
                    document.getElementById('resultado').innerHTML = html;
                }
                setEstado('✅ Listo');
            } catch(e) {
                document.getElementById('resultado').innerHTML = '❌ Error: ' + e.message;
                setEstado('❌ Error');
            }
        }
        
        async function probar() {
            const prov = document.getElementById('proveedor').value;
            const modelo = document.getElementById('modelo').value;
            setEstado('🧪 Probando...');
            try {
                const resp = await fetch('/probar', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({ modelo, proveedor: prov })
                });
                const data = await resp.json();
                document.getElementById('resultado').innerHTML = '✅ ' + (data.respuesta || 'OK');
                setEstado('✅ Modelo responde');
            } catch(e) {
                document.getElementById('resultado').innerHTML = '❌ Error: ' + e.message;
                setEstado('❌ Error');
            }
        }
        
        async function exportar(formato) {
            const contenido = document.getElementById('resultado').innerText;
            if (!contenido || contenido === 'Bienvenido. Escribe un tema y pulsa "Generar Trabajo".') {
                alert('No hay contenido para exportar.');
                return;
            }
            setEstado(`⏳ Exportando a ${formato}...`);
            try {
                const resp = await fetch('/exportar', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({ formato, contenido })
                });
                if (formato === 'txt') {
                    const blob = await resp.blob();
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = 'trabajo.txt';
                    a.click();
                } else {
                    const blob = await resp.blob();
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = 'trabajo.docx';
                    a.click();
                }
                setEstado('✅ Exportado');
            } catch(e) {
                setEstado('❌ Error al exportar');
                alert('Error: ' + e.message);
            }
        }
        
        function limpiar() {
            document.getElementById('resultado').innerHTML = 'Conversación limpiada.';
            setEstado('');
        }
    </script>
    </body>
    </html>
    """
    
    @app.route('/')
    def index():
        return render_template_string(HTML_TEMPLATE)
    
    @app.route('/generar', methods=['POST'])
    def generar():
        data = request.json
        tema = data.get('tema', '').strip()
        modelo = data.get('modelo', 'openai')
        proveedor = data.get('proveedor', 'pollinations')
        if not tema:
            return jsonify({'error': 'Tema vacío'}), 400
        try:
            texto = core.generar_trabajo(tema, modelo, proveedor)
            return jsonify({'texto': texto})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    @app.route('/probar', methods=['POST'])
    def probar():
        data = request.json
        modelo = data.get('modelo', 'openai')
        proveedor = data.get('proveedor', 'pollinations')
        try:
            resp = core.consultar_ia(modelo, [{"role":"user","content":"Responde OK"}], proveedor)
            return jsonify({'respuesta': resp})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    @app.route('/exportar', methods=['POST'])
    def exportar():
        data = request.json
        formato = data.get('formato', 'txt')
        contenido = data.get('contenido', '')
        if not contenido:
            return jsonify({'error': 'Sin contenido'}), 400
        
        if formato == 'txt':
            from io import BytesIO
            buffer = BytesIO()
            buffer.write(contenido.encode('utf-8'))
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name='trabajo.txt', mimetype='text/plain')
        
        elif formato == 'docx':
            if Document is None:
                return jsonify({'error': 'python-docx no instalado'}), 500
            from io import BytesIO
            doc = Document()
            doc.add_heading('Trabajo Académico', 0)
            doc.add_paragraph(f"Generado: {datetime.datetime.now()}")
            for linea in contenido.split('\n'):
                linea = linea.strip()
                if not linea:
                    continue
                if linea.startswith('📌'):
                    doc.add_heading(linea.replace('📌','').strip(), level=1)
                elif linea.startswith('***') or linea.startswith('**'):
                    doc.add_heading(linea.strip('*'), level=2)
                else:
                    doc.add_paragraph(linea)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name='trabajo.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            return jsonify({'error': 'Formato no soportado'}), 400
    
    # Puerto para Render (usar variable de entorno PORT)
    port = int(os.environ.get('PORT', 5000))
    
    if kalm_mode:
        # ═══ MODO KALM: ARRANCAR SERVIDOR EN SEGUNDO PLANO ═══
        print(f"🚀 Iniciando Chat Académico en puerto {port} (modo Kalm)")
        
        # Iniciar el servidor en un hilo separado
        def run_server():
            try:
                app.run(host='0.0.0.0', port=port, debug=False, use_reloader=False)
            except Exception as e:
                print(f"❌ Error en servidor: {e}")
        
        thread = threading.Thread(target=run_server, daemon=True)
        thread.start()
        
        # Esperar 2 segundos a que el servidor esté listo
        time.sleep(2)
        
        # ═══ IMPRIMIR LA URL PARA QUE EL FRONTEND LA CAPTURE ═══
        print(f"http://localhost:{port}")
        sys.stdout.flush()  # Forzar el flush para que el frontend lo capture
        return
    
    else:
        # Modo normal: ejecutar en primer plano
        print(f"🚀 Servidor web iniciado en http://0.0.0.0:{port}")
        app.run(host='0.0.0.0', port=port, debug=False)

# ============================================================
# 6. MODO GUI (Tkinter) - para Windows o con display
# ============================================================
def ejecutar_gui():
    """Arranca la interfaz gráfica Tkinter."""
    try:
        from tkinter import Tk, Frame, Label, Entry, Button, Text, scrolledtext, messagebox, filedialog, ttk
    except ImportError:
        print("❌ Tkinter no disponible. Usando modo web.")
        ejecutar_web(kalm_mode=False)
        return
    
    class ChatAppGUI:
        def __init__(self, root):
            self.root = root
            self.root.title("Chat Académico - IA sin API Key")
            self.root.geometry("900x750")
            self.core = ChatAcademicoCore()
            self.modelo_var = tk.StringVar(value="openai")
            self.proveedor_var = tk.StringVar(value="pollinations")
            self.crear_widgets()
            self.mostrar_bienvenida()
        
        def crear_widgets(self):
            top = ttk.Frame(self.root, padding="10")
            top.pack(fill=tk.X)
            
            ttk.Label(top, text="Proveedor:").pack(side=tk.LEFT, padx=(0,5))
            cb_prov = ttk.Combobox(top, textvariable=self.proveedor_var, values=["pollinations", "g4f"], state="readonly", width=12)
            cb_prov.pack(side=tk.LEFT, padx=(0,10))
            cb_prov.bind("<<ComboboxSelected>>", self.cambiar_proveedor)
            
            ttk.Label(top, text="Modelo:").pack(side=tk.LEFT, padx=(0,5))
            self.combo_modelos = ttk.Combobox(top, textvariable=self.modelo_var, values=["openai","claude","gemini","deepseek","llama","mistral","phi","qwen"], state="readonly", width=20)
            self.combo_modelos.pack(side=tk.LEFT, padx=(0,10))
            
            ttk.Button(top, text="🔍 Probar", command=self.probar_modelo).pack(side=tk.LEFT, padx=(0,10))
            ttk.Label(top, text="Tema:").pack(side=tk.LEFT, padx=(10,5))
            self.entry_tema = ttk.Entry(top, width=40)
            self.entry_tema.pack(side=tk.LEFT, padx=(0,10))
            self.entry_tema.bind("<Return>", lambda e: self.generar_trabajo())
            ttk.Button(top, text="📝 Generar", command=self.generar_trabajo).pack(side=tk.LEFT)
            
            frame_chat = ttk.Frame(self.root, padding="10")
            frame_chat.pack(fill=tk.BOTH, expand=True)
            self.text_area = scrolledtext.ScrolledText(frame_chat, wrap=tk.WORD, font=("Segoe UI", 10))
            self.text_area.pack(fill=tk.BOTH, expand=True)
            self.text_area.tag_configure("usuario", foreground="#0055cc", font=("Segoe UI", 10, "bold"))
            self.text_area.tag_configure("ia", foreground="#008800")
            self.text_area.tag_configure("titulo", foreground="#cc3300", font=("Segoe UI", 12, "bold"))
            self.text_area.tag_configure("subtitulo", foreground="#993300", font=("Segoe UI", 11, "bold"))
            
            bottom = ttk.Frame(self.root, padding="10")
            bottom.pack(fill=tk.X)
            ttk.Button(bottom, text="📤 Exportar Word", command=self.exportar_word).pack(side=tk.LEFT, padx=5)
            ttk.Button(bottom, text="📄 Exportar TXT", command=self.exportar_txt).pack(side=tk.LEFT, padx=5)
            ttk.Button(bottom, text="🗑️ Limpiar", command=self.limpiar_chat).pack(side=tk.LEFT, padx=5)
            ttk.Button(bottom, text="❓ Ayuda", command=self.mostrar_ayuda).pack(side=tk.RIGHT, padx=5)
        
        def mostrar_bienvenida(self):
            self.text_area.insert(tk.END, "🤖 CHAT ACADÉMICO\n━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n✅ Sin API Key (Pollinations AI)\n✅ Elige modelo y tema\n✅ Genera trabajos con Introducción (ES/EN), Desarrollo, Conclusiones y Bibliografía.\n\n", "ia")
        
        def agregar_mensaje(self, texto, tipo="ia"):
            self.text_area.insert(tk.END, texto + "\n\n", tipo)
            self.text_area.see(tk.END)
            self.root.update()
        
        def cambiar_proveedor(self, event=None):
            prov = self.proveedor_var.get()
            if prov == "pollinations":
                self.combo_modelos['values'] = ["openai","claude","gemini","deepseek","llama","mistral","phi","qwen"]
                self.modelo_var.set("openai")
            else:
                self.combo_modelos['values'] = ["gpt-4o-mini","gpt-4o","gpt-3.5-turbo","claude-3-haiku","gemini-pro","llama-3-70b"]
                self.modelo_var.set("gpt-4o-mini")
        
        def probar_modelo(self):
            modelo = self.modelo_var.get()
            prov = self.proveedor_var.get()
            self.agregar_mensaje(f"🧪 Probando {prov}/{modelo}...", "usuario")
            def tarea():
                try:
                    resp = self.core.consultar_ia(modelo, [{"role":"user","content":"Responde OK"}], prov)
                    self.root.after(0, lambda: self.agregar_mensaje(f"✅ Modelo responde: {resp}", "ia"))
                except Exception as e:
                    self.root.after(0, lambda: self.agregar_mensaje(f"❌ Error: {e}", "ia"))
            threading.Thread(target=tarea, daemon=True).start()
        
        def generar_trabajo(self):
            tema = self.entry_tema.get().strip()
            if not tema:
                messagebox.showwarning("Tema vacío", "Escribe un tema.")
                return
            modelo = self.modelo_var.get()
            prov = self.proveedor_var.get()
            self.agregar_mensaje(f"📝 Generando sobre: {tema}", "usuario")
            def tarea():
                try:
                    texto = self.core.generar_trabajo(tema, modelo, prov)
                    self.root.after(0, lambda: self.agregar_mensaje(texto, "ia"))
                except Exception as e:
                    self.root.after(0, lambda: self.agregar_mensaje(f"❌ Error: {e}", "ia"))
            threading.Thread(target=tarea, daemon=True).start()
        
        def limpiar_chat(self):
            self.text_area.delete(1.0, tk.END)
            self.mostrar_bienvenida()
        
        def obtener_contenido(self):
            return self.text_area.get(1.0, tk.END).strip()
        
        def exportar_txt(self):
            contenido = self.obtener_contenido()
            if not contenido:
                messagebox.showinfo("Sin contenido", "No hay nada para exportar.")
                return
            archivo = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("TXT","*.txt")])
            if archivo:
                with open(archivo, 'w', encoding='utf-8') as f:
                    f.write(contenido)
                messagebox.showinfo("Éxito", f"Guardado en {archivo}")
        
        def exportar_word(self):
            if Document is None:
                messagebox.showerror("Error", "Instala python-docx: pip install python-docx")
                return
            contenido = self.obtener_contenido()
            if not contenido:
                messagebox.showinfo("Sin contenido", "No hay nada para exportar.")
                return
            archivo = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word","*.docx")])
            if archivo:
                doc = Document()
                doc.add_heading('Trabajo Académico - Chat IA', 0)
                doc.add_paragraph(f"Generado: {datetime.datetime.now()}")
                for linea in contenido.split('\n'):
                    linea = linea.strip()
                    if not linea:
                        continue
                    if linea.startswith('📌'):
                        doc.add_heading(linea.replace('📌','').strip(), level=1)
                    elif linea.startswith('***') or linea.startswith('**'):
                        doc.add_heading(linea.strip('*'), level=2)
                    else:
                        doc.add_paragraph(linea)
                doc.save(archivo)
                messagebox.showinfo("Éxito", f"Guardado en {archivo}")
        
        def mostrar_ayuda(self):
            messagebox.showinfo("Ayuda", "Usa Pollinations AI (gratis, sin key).\nModelos: openai, claude, gemini, etc.\nGenera trabajos con estructura académica.")
    
    root = Tk()
    app = ChatAppGUI(root)
    root.mainloop()

# ============================================================
# 7. PUNTO DE ENTRADA PRINCIPAL
# ============================================================
if __name__ == "__main__":
    # Verificar si se ejecuta en modo Kalm
    kalm_mode = "--kalm" in sys.argv
    
    print("🔍 Detectando entorno...")
    if es_headless() or kalm_mode:
        print("🌐 Entorno headless detectado. Arrancando servidor web (Flask).")
        ejecutar_web(kalm_mode=kalm_mode)
    else:
        print("🖥️ Entorno gráfico detectado. Arrancando GUI (Tkinter).")
        ejecutar_gui()